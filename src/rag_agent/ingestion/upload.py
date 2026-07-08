"""通用檔案上傳 → 知識庫 ingest：解析文字 → 切段 → embedding → 新 source/chunks。

供設定頁「知識庫管理」的管理者上傳文件。支援 TXT / MD / PDF / DOCX。
知識庫（Source）可動態建立（source_type="uploaded"），與 connector 式來源共用檢索/證據流程。
"""

import hashlib
import io
import uuid
from datetime import datetime
from pathlib import Path

from sqlalchemy import text as sqltext
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import select

from ..config import settings
from ..db.models import Chunk, RawDocument, Source
from ..governance.chunking import _sliding_windows, _split_paragraphs
from ..indexing.embedding import embed_pending

SUPPORTED = (".txt", ".md", ".pdf", ".docx")


def _extract_text(filename: str, raw: bytes) -> str:
    """依副檔名擷取純文字。"""
    ext = Path(filename).suffix.lower()
    if ext in (".txt", ".md"):
        return raw.decode("utf-8", errors="replace")
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        import docx
        d = docx.Document(io.BytesIO(raw))
        return "\n".join(p.text for p in d.paragraphs)
    raise ValueError(f"不支援的檔案格式：{ext}（支援 {', '.join(SUPPORTED)}）")


async def create_kb(session: AsyncSession, name: str) -> Source:
    """建立一個新的（空的）使用者知識庫。"""
    src = Source(
        source_id="kb_" + uuid.uuid4().hex[:8],
        source_name=name,
        publisher="使用者上傳",
        source_type="uploaded",
        jurisdiction="TW",
        license_type="user_provided",
        access_method="UPLOAD",
        update_frequency="on_demand",
        trust_score=70,
        attribution_required=False,
        full_text_indexing=True,
        phase="Phase3",
        enabled=True,
    )
    session.add(src)
    await session.commit()
    await session.refresh(src)
    return src


async def delete_kb(session: AsyncSession, source_id: str) -> None:
    """刪除知識庫及其所有文件與 chunk。"""
    await session.execute(sqltext("DELETE FROM chunks WHERE source_id = :s"), {"s": source_id})
    await session.execute(
        sqltext("DELETE FROM raw_documents WHERE source_id = :s"), {"s": source_id}
    )
    await session.execute(sqltext("DELETE FROM sources WHERE source_id = :s"), {"s": source_id})
    await session.commit()


async def ingest_file(
    session: AsyncSession,
    source_id: str,
    filename: str,
    raw: bytes,
    chunk_size: int = 512,
    chunk_overlap: int = 64,
    credibility: int = 70,
) -> dict:
    """把上傳檔案解析 → 切段 → 寫入 chunks → embedding。回傳統計。"""
    checksum = f"sha256:{hashlib.sha256(raw).hexdigest()}"
    dup = await session.execute(select(RawDocument).where(RawDocument.checksum == checksum))
    if dup.first():
        return {"skipped": True, "reason": "duplicate", "filename": filename}

    content = _extract_text(filename, raw)
    if not content.strip():
        raise ValueError("無法從檔案擷取到文字內容")

    ext = Path(filename).suffix.lower()
    dest = settings.raw_dir / source_id / (checksum.removeprefix("sha256:") + ext)
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(raw)

    doc = RawDocument(
        source_id=source_id,
        source_type="uploaded",
        source_url="",
        raw_format=ext.lstrip("."),
        content_pointer=str(dest),
        fetched_at=datetime.utcnow(),
        source_version=filename,
        checksum=checksum,
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    windows = _sliding_windows(_split_paragraphs(content), chunk_size, chunk_overlap)
    for i, w in enumerate(windows):
        session.add(Chunk(
            chunk_id=f"UP-{source_id}-{doc.id}-p{i}",
            raw_document_id=doc.id,
            source_id=source_id,
            source_type="uploaded",
            title=filename,
            text=w,
            section_path=f"p{i}",
            original_url="",
            published_at=None,
            credibility_score=credibility,
            chunk_index=i,
            token_count=max(1, len(w) // 2),
            checksum=f"sha256:{hashlib.sha256(w.encode()).hexdigest()}",
        ))
    await session.commit()

    embedded = await embed_pending(session)
    return {
        "doc_id": doc.id,
        "filename": filename,
        "chunks_added": len(windows),
        "chunks_embedded": embedded,
    }


async def list_documents(session: AsyncSession, source_id: str) -> list[dict]:
    """列出某知識庫的所有文件（含 chunk 數）。"""
    rows = await session.execute(sqltext("""
        SELECT r.id, r.source_version AS filename, r.raw_format, r.fetched_at,
               COUNT(c.chunk_id) AS chunk_count
        FROM raw_documents r
        LEFT JOIN chunks c ON c.raw_document_id = r.id
        WHERE r.source_id = :s
        GROUP BY r.id
        ORDER BY r.fetched_at DESC
    """), {"s": source_id})
    out = []
    for r in rows.mappings():
        d = dict(r)
        d["fetched_at"] = d["fetched_at"].isoformat() if d["fetched_at"] else ""
        out.append(d)
    return out


async def delete_document(session: AsyncSession, doc_id: int) -> None:
    """刪除單一文件及其 chunk。"""
    await session.execute(sqltext("DELETE FROM chunks WHERE raw_document_id = :i"), {"i": doc_id})
    await session.execute(sqltext("DELETE FROM raw_documents WHERE id = :i"), {"i": doc_id})
    await session.commit()
